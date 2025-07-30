import os
import time
import pdfplumber
from docx import Document
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import subprocess
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

ENABLE_OLLAMA = True  # Set to False to disable Ollama LLM translation
ENABLE_NLLB = False  # Set to False to disable NLLB translation

# Validate flags
if not ENABLE_OLLAMA and not ENABLE_NLLB:
    print("Error: At least one translation method must be enabled!")
    print("Set ENABLE_OLLAMA=True and/or ENABLE_NLLB=True")
    exit(1)

print(f"Translation methods enabled:")
print(f"  NLLB: {'Yes' if ENABLE_NLLB else 'No'}")
print(f"  Ollama: {'Yes' if ENABLE_OLLAMA else 'No'}")
print("-" * 40)

def translate_with_context(text: str, context_paragraphs: list = None, max_retries: int = 3) -> str:
    """
    Translate text with optional context from previous paragraphs
    """
    for attempt in range(max_retries + 1):
        # Build context section
        context_section = ""
        if context_paragraphs and len(context_paragraphs) > 0:
            context_section = "\n\nPrevious context for better translation:\n"
            for i, ctx in enumerate(context_paragraphs[-3:], 1):  # Use last 3 paragraphs as context
                context_section += f"Context {i}: \"{ctx}\"\n"
        
        prompt = f"""
You are a professional translation machine. Translate the following sentence from Slovak to English.
Use the provided context to ensure consistency and accuracy.
Respond ONLY in this format:
translated: <translated text>
{context_section}
Current sentence to translate: "{text}"
translated:""".strip()

        # Run the llama3.1 model through Ollama
        result = subprocess.run(
            ["ollama", "run", "llama3.1:latest", prompt],
            capture_output=True,
            text=True
        )

        # Extract output
        output = result.stdout.strip()

        # Find the line that starts with "translated:"
        for line in output.splitlines():
            if line.lower().startswith("translated:"):
                # Clean up the translation
                translation = line[11:].strip()  # Remove "translated:" prefix
                
                # Remove quotes if present
                if translation.startswith('"') and translation.endswith('"'):
                    translation = translation[1:-1].strip()
                elif translation.startswith("'") and translation.endswith("'"):
                    translation = translation[1:-1].strip()
                
                return translation

        # If no valid translation found and we have retries left
        if attempt < max_retries:
            print(f"  Retry attempt {attempt + 1}/{max_retries} for translation...")
            time.sleep(1)  # Brief pause before retry
        
    return f"[TRANSLATION FAILED AFTER {max_retries} RETRIES] {text}"

# Add formatting function
def apply_translation_with_formatting(paragraph, translated_text, original_text, error_prefix=""):
    """Apply translation while preserving formatting, with error handling."""
    try:
        # Collect all formatting from existing runs
        original_runs_formatting = []
        for run in paragraph.runs:
            formatting = {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'color': run.font.color.rgb if run.font.color.rgb else None
            }
            original_runs_formatting.append(formatting)
        
        # Clear paragraph and add translated text
        paragraph.clear()
        new_run = paragraph.add_run(translated_text)
        
        # Apply formatting from first run if available
        if original_runs_formatting:
            fmt = original_runs_formatting[0]
            if fmt['bold'] is not None:
                new_run.bold = fmt['bold']
            if fmt['italic'] is not None:
                new_run.italic = fmt['italic']
            if fmt['underline'] is not None:
                new_run.underline = fmt['underline']
            if fmt['font_name']:
                new_run.font.name = fmt['font_name']
            if fmt['font_size']:
                new_run.font.size = fmt['font_size']
            if fmt['color']:
                new_run.font.color.rgb = fmt['color']
        
        return True
        
    except Exception as e:
        print(f"ERROR: Could not apply formatting: {e}")
        return False

# Replace the old translate function call with the new one
def translate(text: str, max_retries: int = 3) -> str:
    """Backward compatibility wrapper"""
    return translate_with_context(text, None, max_retries)

# Initialize NLLB model only if enabled
if ENABLE_NLLB:
    # "facebook/nllb-200-distilled-600M" is a smaller model, but "facebook/nllb-200-distilled-1.3B" is more accurate
    model_name = "facebook/nllb-200-distilled-600M"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)

    src_text = "Hello, how are you?"
    tokenizer.src_lang = "eng_Latn"
    tokens = tokenizer(src_text, return_tensors="pt")
    generated = model.generate(**tokens, forced_bos_token_id=tokenizer.convert_tokens_to_ids("slk_Latn"))
    print(tokenizer.decode(generated[0], skip_special_tokens=True))

def extract_text_from_docx(docx_path):
    results = []
    doc = Document(docx_path)
    
    for para_num, paragraph in enumerate(doc.paragraphs, start=1):
        if paragraph.text.strip():  
            results.append({
                "text": paragraph.text.strip(),
                "bbox": None,  
                "page": para_num 
            })
    return results

def extract_text_with_layout(pdf_path, merge_threshold=10):
    results = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            words = page.extract_words()
            if not words:
                continue
            
            # Group words into lines/phrases
            current_group = {
                "text": words[0]["text"],
                "bbox": (words[0]["x0"], words[0]["top"], words[0]["x1"], words[0]["bottom"]),
                "page": page_num,
                "words": [words[0]]
            }
            
            for word in words[1:]:
                # Check if word is close enough to current group
                prev_word = current_group["words"][-1]
                
                # Merge if words are on same line (similar y-coordinates) and close horizontally
                if (abs(word["top"] - prev_word["top"]) <= merge_threshold and 
                    word["x0"] - prev_word["x1"] <= merge_threshold * 2):
                    
                    # Extend the group
                    current_group["text"] += " " + word["text"]
                    current_group["bbox"] = (
                        min(current_group["bbox"][0], word["x0"]),
                        min(current_group["bbox"][1], word["top"]),
                        max(current_group["bbox"][2], word["x1"]),
                        max(current_group["bbox"][3], word["bottom"])
                    )
                    current_group["words"].append(word)
                else:
                    # Save current group and start new one
                    results.append({
                        "text": current_group["text"],
                        "bbox": current_group["bbox"],
                        "page": current_group["page"]
                    })
                    
                    current_group = {
                        "text": word["text"],
                        "bbox": (word["x0"], word["top"], word["x1"], word["bottom"]),
                        "page": page_num,
                        "words": [word]
                    }
            
            if current_group:
                results.append({
                    "text": current_group["text"],
                    "bbox": current_group["bbox"],
                    "page": current_group["page"]
                })
    
    return results


docx_file = "User manual ProfileManagerWeb_v4.3.301 ENG.docx"  # Your .docx file

# === NLLB TRANSLATION ===
if ENABLE_NLLB:
    print("\n" + "="*60)
    print("Starting NLLB Translation...")
    print("="*60)
    
    doc = Document(docx_file)

    count = 0
    total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
    start_time = time.time()

    print(f"Starting translation of {total_paragraphs} paragraphs...")

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Calculate progress and time estimates
            current_time = time.time()
            elapsed_time = current_time - start_time
            
            if count > 0:
                avg_time_per_paragraph = elapsed_time / count
                remaining_paragraphs = total_paragraphs - count
                estimated_remaining_time = remaining_paragraphs * avg_time_per_paragraph
                
                # Format time estimates
                elapsed_mins = int(elapsed_time // 60)
                elapsed_secs = int(elapsed_time % 60)
                remaining_mins = int(estimated_remaining_time // 60)
                remaining_secs = int(estimated_remaining_time % 60)
                
                print(f"Progress {count}/{total_paragraphs} | Elapsed: {elapsed_mins:02d}:{elapsed_secs:02d} | Est. remaining: {remaining_mins:02d}:{remaining_secs:02d}")
            else:
                print(f"Progress {count}/{total_paragraphs} | Starting...")
            
            count += 1
            
            original_text = paragraph.text.strip()
            print(f"[Para {count}] {original_text}")
            
            # Translate the text from Slovak to English
            tokenizer.src_lang = "slk_Latn"  # Set source language to Slovak
            tokens = tokenizer(original_text, return_tensors="pt")
            generated = model.generate(**tokens, forced_bos_token_id=tokenizer.convert_tokens_to_ids("eng_Latn"))
            translated_text = tokenizer.decode(generated[0], skip_special_tokens=True)
            
            # Attempt to apply original formatting
            success = apply_translation_with_formatting(paragraph, translated_text, original_text, "NLLB ")
            if success:
                print(f"Translated: {translated_text}")
            else:
                print(f"Keeping original text and highlighting it")
                
                # Keep original text but highlight it for manual review
                paragraph.clear()
                highlighted_run = paragraph.add_run(f"[TRANSLATION ERROR - MANUAL REVIEW NEEDED] {original_text}")
                
                # Add yellow highlighting if possible
                try:
                    highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    highlighted_run.font.color.rgb = RGBColor(255, 0, 0)  # Red text
                    highlighted_run.bold = True
                except:
                    # If highlighting fails, just make it bold and add marker
                    highlighted_run.bold = True
            
            print("-" * 40)


    # Final time summary for NLLB version
    total_time = time.time() - start_time
    total_mins = int(total_time // 60)
    total_secs = int(total_time % 60)

    print(f"\nNLLB Translation completed!")
    print(f"Total paragraphs translated: {total_paragraphs}")
    print(f"Total time: {total_mins:02d}:{total_secs:02d}")
    print(f"Average time per paragraph: {total_time/total_paragraphs:.2f} seconds")

    # Save the NLLB translated document
    output_filename = "User_manual_ProfileManagerWeb_v4.3.301_ENG_NLLB.docx"
    doc.save(output_filename)
    print(f"NLLB translated document saved as: {output_filename}")

# === OLLAMA LLM TRANSLATION ===
if ENABLE_OLLAMA:
    print("\n" + "="*60)
    print("Starting OLLAMA LLM Translation...")
    print("="*60)

    # Load fresh document for Ollama translation
    doc_ollama = Document(docx_file)

    count_ollama = 0
    total_paragraphs_ollama = len([p for p in doc_ollama.paragraphs if p.text.strip()])
    start_time_ollama = time.time()
    
    # Store context for better translation
    translated_context = []  # Keep track of translated paragraphs for context

    print(f"Starting Ollama translation of {total_paragraphs_ollama} paragraphs...")

    for paragraph in doc_ollama.paragraphs:
        if paragraph.text.strip():
            # Calculate progress and time estimates for Olloma
            current_time = time.time()
            elapsed_time = current_time - start_time_ollama
            
            if count_ollama > 0:
                avg_time_per_paragraph = elapsed_time / count_ollama
                remaining_paragraphs = total_paragraphs_ollama - count_ollama
                estimated_remaining_time = remaining_paragraphs * avg_time_per_paragraph
                
                # Format time estimates
                elapsed_mins = int(elapsed_time // 60)
                elapsed_secs = int(elapsed_time % 60)
                remaining_mins = int(estimated_remaining_time // 60)
                remaining_secs = int(estimated_remaining_time % 60)
                
                print(f"Ollama Progress {count_ollama}/{total_paragraphs_ollama} | Elapsed: {elapsed_mins:02d}:{elapsed_secs:02d} | Est. remaining: {remaining_mins:02d}:{remaining_secs:02d}")
            else:
                print(f"Ollama Progress {count_ollama}/{total_paragraphs_ollama} | Starting...")
            
            count_ollama += 1
            
            original_text = paragraph.text.strip()
            print(f"[Ollama Para {count_ollama}] {original_text}")
            
            # Use Ollama LLM translation with context
            try:
                # Use context-aware translation
                translated_text = translate_with_context(original_text, translated_context)
                
                # Remove the "translated:" prefix if present
                if translated_text.lower().startswith("translated:"):
                    translated_text = translated_text[11:].strip()
                
                # Remove quotes at the start and end
                if translated_text.startswith('"') and translated_text.endswith('"'):
                    translated_text = translated_text[1:-1].strip()
                elif translated_text.startswith("'") and translated_text.endswith("'"):
                    translated_text = translated_text[1:-1].strip()
                
                # Check if translation failed after retries
                if translated_text.startswith("[TRANSLATION FAILED AFTER"):
                    print(f"Ollama Translation FAILED: {translated_text}")
                    # Apply highlighting for failed translation
                    paragraph.clear()
                    highlighted_run = paragraph.add_run(translated_text)
                    
                    # Add red highlighting for failed translations
                    try:
                        highlighted_run.font.highlight_color = WD_COLOR_INDEX.RED
                        highlighted_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                        highlighted_run.bold = True
                    except:
                        highlighted_run.bold = True
                    continue  # Skip to next paragraph
                
                # Add successful translation to context for next paragraphs
                translated_context.append(translated_text)
                # Keep only last 5 paragraphs as context to avoid too much text
                if len(translated_context) > 5:
                    translated_context.pop(0)
                
                print(f"Ollama Translated: {translated_text}")
            except Exception as e:
                print(f"WARNING: Ollama translation failed: {e}")
                translated_text = f"[OLLAMA TRANSLATION FAILED] {original_text}"
                
                # Apply highlighting for exception-based failures
                paragraph.clear()
                highlighted_run = paragraph.add_run(translated_text)
                
                try:
                    highlighted_run.font.highlight_color = WD_COLOR_INDEX.RED
                    highlighted_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                    highlighted_run.bold = True
                except:
                    highlighted_run.bold = True
                continue  # Skip to next paragraph
            
            # Apply translation with error handling
            success = apply_translation_with_formatting(paragraph, translated_text, original_text, "OLLAMA ")
            if success:
                print(f"Ollama Applied: {translated_text}")
            else:
                print(f"Keeping original text and highlighting it")
                
                # Keep original text but highlight it for manual review
                paragraph.clear()
                highlighted_run = paragraph.add_run(f"[OLLAMA FORMATTING ERROR - MANUAL REVIEW NEEDED] {original_text}")
                
                # Add yellow highlighting if possible
                try:
                    highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    highlighted_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue text for Olloma errors
                    highlighted_run.bold = True
                except:
                    # If highlighting fails, just make it bold and add marker
                    highlighted_run.bold = True
            
            print("-" * 40)

    # Final time summary for Ollama version
    total_time_ollama = time.time() - start_time_ollama
    total_mins_ollama = int(total_time_ollama // 60)
    total_secs_ollama = int(total_time_ollama % 60)

    print(f"\nOllama Translation completed!")
    print(f"Total paragraphs translated: {total_paragraphs_ollama}")
    print(f"Total time: {total_mins_ollama:02d}:{total_secs_ollama:02d}")
    print(f"Average time per paragraph: {total_time_ollama/total_paragraphs_ollama:.2f} seconds")

    # Save the Ollama translated document
    import time

    readable_time = time.strftime("%Y%m%d_%H%M%S", time.localtime(current_time))

    output_filename_ollama = f"User_manual_ProfileManagerWeb_v4.3.301_ENG_OLLAMA_{readable_time}.docx"
    doc_ollama.save(output_filename_ollama)
    print(f"Olloma translated document saved as: {output_filename_ollama}")

# Final summary
print("\n" + "="*60)
print("TRANSLATION COMPLETED!")
if ENABLE_NLLB:
    print(f"NLLB Version: User_manual_ProfileManagerWeb_v4.3.301_ENG_NLLB.docx")
if ENABLE_OLLAMA:
    print(f"Olloma Version: User_manual_ProfileManagerWeb_v4.3.301_ENG_OLLAMA.docx")
print("="*60)
